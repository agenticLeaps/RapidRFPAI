"""
AWS Bedrock Client for Claude AI
Used for workspace AI and extraction agents
"""

import os
import json
import time
import base64
from typing import List, Tuple, Dict, Any, Optional

# Bedrock availability flag
BEDROCK_AVAILABLE = False
bedrock_client = None

try:
    import boto3
    from botocore.exceptions import ClientError
    from botocore.config import Config
    BEDROCK_AVAILABLE = True
except ImportError:
    print("⚠️ boto3 not installed. Bedrock features will be unavailable.")


def init_bedrock_client():
    """Initialize Bedrock client with AWS credentials"""
    global bedrock_client, BEDROCK_AVAILABLE

    if not BEDROCK_AVAILABLE:
        return None

    try:
        region = os.environ.get("AWS_REGION", "us-east-1")

        # Configure longer timeout for large Claude extractions (5 minutes)
        bedrock_config = Config(
            read_timeout=300,
            connect_timeout=30,
            retries={'max_attempts': 3}
        )

        # Build client kwargs - only include credentials if explicitly set
        # This allows boto3 to use ~/.aws/credentials file as fallback
        client_kwargs = {
            'service_name': 'bedrock-runtime',
            'region_name': region,
            'config': bedrock_config
        }

        # Only pass credentials if environment variables are set
        access_key = os.environ.get("AWS_ACCESS_KEY_ID")
        secret_key = os.environ.get("AWS_SECRET_ACCESS_KEY")
        if access_key and secret_key:
            client_kwargs['aws_access_key_id'] = access_key
            client_kwargs['aws_secret_access_key'] = secret_key

        bedrock_client = boto3.client(**client_kwargs)
        print(f"✅ Bedrock client initialized (region: {region})")
        return bedrock_client
    except Exception as e:
        print(f"❌ Failed to initialize Bedrock client: {e}")
        BEDROCK_AVAILABLE = False
        return None


def get_bedrock_client():
    """Get or create Bedrock client"""
    global bedrock_client
    if bedrock_client is None:
        init_bedrock_client()
    return bedrock_client


class BedrockClaude:
    """Claude AI client using AWS Bedrock"""

    # Model IDs
    CLAUDE_SONNET = "us.anthropic.claude-sonnet-4-5-20250929-v1:0"
    CLAUDE_HAIKU = "us.anthropic.claude-3-haiku-20240307-v1:0"

    def __init__(self, model_id: str = None):
        self.client = get_bedrock_client()
        self.model_id = model_id or os.environ.get("BEDROCK_MODEL_ID", self.CLAUDE_SONNET)
        self.max_retries = 3
        self.retry_delay = 5  # seconds

    def call_claude(
        self,
        prompt: str,
        system: str = None,
        max_tokens: int = 4096,
        temperature: float = 0.2,
        response_format: str = "json"
    ) -> Dict[str, Any]:
        """
        Call Claude via Bedrock

        Args:
            prompt: User message/prompt
            system: System prompt (optional)
            max_tokens: Maximum output tokens
            temperature: Response temperature (0-1)
            response_format: "json" to parse as JSON, "text" for raw text

        Returns:
            Parsed JSON dict or {"text": response} for text format
        """
        if not self.client:
            raise RuntimeError("Bedrock client not initialized. Check AWS credentials.")

        messages = [{"role": "user", "content": prompt}]

        # Build request body
        body = {
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": max_tokens,
            "temperature": temperature,
            "messages": messages
        }

        if system:
            body["system"] = system

        # Retry logic with exponential backoff
        last_error = None
        for attempt in range(self.max_retries):
            try:
                response = self.client.invoke_model(
                    modelId=self.model_id,
                    body=json.dumps(body)
                )

                response_body = json.loads(response['body'].read())
                response_text = response_body['content'][0]['text'].strip()

                # Parse JSON if requested
                if response_format == "json":
                    return self._parse_json_response(response_text)
                else:
                    return {"text": response_text}

            except ClientError as e:
                error_code = e.response.get('Error', {}).get('Code', '')
                if error_code in ['ThrottlingException', 'ServiceUnavailableException']:
                    last_error = e
                    if attempt < self.max_retries - 1:
                        wait_time = self.retry_delay * (2 ** attempt)
                        print(f"⚠️ Bedrock rate limited (attempt {attempt + 1}/{self.max_retries}). Waiting {wait_time}s...")
                        time.sleep(wait_time)
                    else:
                        raise
                else:
                    raise
            except Exception as e:
                last_error = e
                if attempt < self.max_retries - 1:
                    wait_time = self.retry_delay * (2 ** attempt)
                    print(f"⚠️ Bedrock error (attempt {attempt + 1}/{self.max_retries}): {e}. Retrying in {wait_time}s...")
                    time.sleep(wait_time)
                else:
                    raise

        if last_error:
            raise last_error

    def call_claude_with_images(
        self,
        prompt: str,
        images: List[Tuple[bytes, str]],
        system: str = None,
        max_tokens: int = 4096,
        temperature: float = 0.2,
        response_format: str = "json"
    ) -> Dict[str, Any]:
        """
        Call Claude with image attachments

        Args:
            prompt: User message/prompt
            images: List of (image_bytes, media_type) tuples
            system: System prompt (optional)
            max_tokens: Maximum output tokens
            temperature: Response temperature
            response_format: "json" or "text"

        Returns:
            Parsed response
        """
        if not self.client:
            raise RuntimeError("Bedrock client not initialized. Check AWS credentials.")

        # Build content array with images and text
        content = []

        for img_bytes, media_type in images:
            content.append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": media_type,
                    "data": base64.b64encode(img_bytes).decode('utf-8')
                }
            })

        content.append({"type": "text", "text": prompt})

        messages = [{"role": "user", "content": content}]

        body = {
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": max_tokens,
            "temperature": temperature,
            "messages": messages
        }

        if system:
            body["system"] = system

        # Retry logic
        last_error = None
        for attempt in range(self.max_retries):
            try:
                response = self.client.invoke_model(
                    modelId=self.model_id,
                    body=json.dumps(body)
                )

                response_body = json.loads(response['body'].read())
                response_text = response_body['content'][0]['text'].strip()

                if response_format == "json":
                    return self._parse_json_response(response_text)
                else:
                    return {"text": response_text}

            except Exception as e:
                last_error = e
                if attempt < self.max_retries - 1:
                    wait_time = self.retry_delay * (2 ** attempt)
                    print(f"⚠️ Bedrock error (attempt {attempt + 1}/{self.max_retries}): {e}. Retrying in {wait_time}s...")
                    time.sleep(wait_time)
                else:
                    raise

        if last_error:
            raise last_error

    def call_claude_with_documents(
        self,
        prompt: str,
        documents: List[Tuple[bytes, str]],
        system: str = None,
        max_tokens: int = 32768,
        temperature: float = 0.2,
        response_format: str = "json"
    ) -> Dict[str, Any]:
        """
        Call Claude with document content (PDFs converted to text, images as base64)

        Args:
            prompt: User message/prompt
            documents: List of (file_bytes, filename) tuples
            system: System prompt (optional)
            max_tokens: Maximum output tokens
            temperature: Response temperature
            response_format: "json" or "text"

        Returns:
            Parsed response
        """
        if not self.client:
            raise RuntimeError("Bedrock client not initialized. Check AWS credentials.")

        content = []

        # Process each document
        for file_bytes, filename in documents:
            ext = filename.lower().split('.')[-1]

            # Handle images directly
            if ext in ['png', 'jpg', 'jpeg', 'gif', 'webp']:
                media_type = f"image/{ext}" if ext != 'jpg' else "image/jpeg"
                content.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": media_type,
                        "data": base64.b64encode(file_bytes).decode('utf-8')
                    }
                })
                content.append({"type": "text", "text": f"\n--- Image: {filename} ---\n"})

            # Handle PDFs - extract text
            elif ext == 'pdf':
                try:
                    from PyPDF2 import PdfReader
                    import io
                    reader = PdfReader(io.BytesIO(file_bytes))
                    text_parts = []
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    pdf_text = "\n\n".join(text_parts)
                    content.append({"type": "text", "text": f"\n--- Document: {filename} ---\n\n{pdf_text}\n\n--- End of {filename} ---\n"})
                except Exception as e:
                    print(f"⚠️ Failed to extract PDF text from {filename}: {e}")
                    content.append({"type": "text", "text": f"\n--- Document: {filename} (PDF extraction failed) ---\n"})

            # Handle DOCX
            elif ext == 'docx':
                try:
                    from docx import Document as DocxDocument
                    import io
                    doc = DocxDocument(io.BytesIO(file_bytes))
                    text_parts = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text)
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                            if row_text:
                                text_parts.append(" | ".join(row_text))
                    docx_text = "\n\n".join(text_parts)
                    content.append({"type": "text", "text": f"\n--- Document: {filename} ---\n\n{docx_text}\n\n--- End of {filename} ---\n"})
                except Exception as e:
                    print(f"⚠️ Failed to extract DOCX text from {filename}: {e}")
                    content.append({"type": "text", "text": f"\n--- Document: {filename} (DOCX extraction failed) ---\n"})

            # Handle text files
            elif ext in ['txt', 'csv', 'json']:
                try:
                    text = file_bytes.decode('utf-8')
                    content.append({"type": "text", "text": f"\n--- Document: {filename} ---\n\n{text}\n\n--- End of {filename} ---\n"})
                except:
                    content.append({"type": "text", "text": f"\n--- Document: {filename} (decode failed) ---\n"})

            else:
                content.append({"type": "text", "text": f"\n--- Document: {filename} (unsupported format: {ext}) ---\n"})

        # Add the prompt at the end
        content.append({"type": "text", "text": f"\n\n{prompt}"})

        messages = [{"role": "user", "content": content}]

        body = {
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": max_tokens,
            "temperature": temperature,
            "messages": messages
        }

        if system:
            body["system"] = system

        # Retry logic
        last_error = None
        for attempt in range(self.max_retries):
            try:
                response = self.client.invoke_model(
                    modelId=self.model_id,
                    body=json.dumps(body)
                )

                response_body = json.loads(response['body'].read())
                response_text = response_body['content'][0]['text'].strip()

                print(f"✅ Claude response received ({len(response_text)} chars)")

                if response_format == "json":
                    return self._parse_json_response(response_text)
                else:
                    return {"text": response_text}

            except Exception as e:
                last_error = e
                if attempt < self.max_retries - 1:
                    wait_time = self.retry_delay * (2 ** attempt)
                    print(f"⚠️ Bedrock error (attempt {attempt + 1}/{self.max_retries}): {e}. Retrying in {wait_time}s...")
                    time.sleep(wait_time)
                else:
                    raise

        if last_error:
            raise last_error

    def _parse_json_response(self, response_text: str) -> Dict[str, Any]:
        """Parse JSON from Claude response, handling markdown code blocks"""
        text = response_text.strip()

        # Remove markdown code blocks
        if text.startswith('```json'):
            text = text[7:]
        if text.startswith('```'):
            text = text[3:]
        if text.endswith('```'):
            text = text[:-3]
        text = text.strip()

        try:
            return json.loads(text)
        except json.JSONDecodeError as e:
            print(f"⚠️ JSON parse error: {e}")
            # Try to repair truncated JSON
            return self._repair_json(text)

    def _repair_json(self, text: str) -> Dict[str, Any]:
        """Attempt to repair truncated or malformed JSON"""
        # Count brackets
        open_braces = text.count('{') - text.count('}')
        open_brackets = text.count('[') - text.count(']')

        # Close unclosed structures
        repaired = text.rstrip()

        # Remove trailing incomplete content
        while repaired and repaired[-1] not in '"}]':
            repaired = repaired[:-1]

        repaired += ']' * max(0, open_brackets)
        repaired += '}' * max(0, open_braces)

        try:
            return json.loads(repaired)
        except json.JSONDecodeError:
            print(f"❌ JSON repair failed")
            return {"error": "Failed to parse response", "raw": text[:500]}


# Create singleton instance
claude = BedrockClaude()
