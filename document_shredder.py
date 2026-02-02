"""
Document Shredding Service for V3 Project Creation
Extracts metadata and submission requirements from RFP documents using Gemini
"""

import os
import json
import tempfile
from typing import List, Dict, Any, Optional
from datetime import datetime
from google.cloud import storage
from google.oauth2 import service_account
from vertexai.generative_models import GenerativeModel, Part, SafetySetting

# For docx conversion
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not installed. DOCX files will need conversion.")


# Initialize GCS client with proper credential handling
def get_gcs_client():
    """Get GCS client with proper credential handling"""
    cred_path = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS", "fire.json")

    if cred_path.startswith("{"):
        # JSON string from environment - parse and use credentials
        cred_dict = json.loads(cred_path)
        credentials = service_account.Credentials.from_service_account_info(cred_dict)
        client = storage.Client(credentials=credentials, project=os.environ.get("GOOGLE_CLOUD_PROJECT"))
        print("✅ Using environment credentials for GCS bucket access")
    else:
        # Fallback to file path or default credentials
        if os.path.exists(cred_path):
            credentials = service_account.Credentials.from_service_account_file(cred_path)
            client = storage.Client(credentials=credentials)
            print(f"✅ Using credentials from {cred_path}")
        else:
            client = storage.Client()
            print("✅ Using default credentials for GCS bucket access")

    return client


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text content from a .docx file

    Args:
        file_bytes: The raw bytes of the .docx file

    Returns:
        Extracted text content
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required to process .docx files. Install with: pip install python-docx")

    import io
    doc = DocxDocument(io.BytesIO(file_bytes))

    text_parts = []

    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))

    return "\n\n".join(text_parts)


def download_file_from_gcs(gcs_url: str, temp_dir: str) -> tuple[str, str]:
    """
    Download a file from Google Cloud Storage to a temporary directory

    Args:
        gcs_url: GCS URL (gs://bucket/path/to/file)
        temp_dir: Temporary directory to save the file

    Returns:
        Tuple of (local_file_path, filename)
    """
    if not gcs_url.startswith('gs://'):
        raise ValueError(f"Invalid GCS URL: {gcs_url}")

    # Parse GCS URL
    parts = gcs_url.replace('gs://', '').split('/', 1)
    bucket_name = parts[0]
    file_path = parts[1] if len(parts) > 1 else ''

    # Download file using GCS client
    gcs_client = get_gcs_client()
    bucket = gcs_client.bucket(bucket_name)
    blob = bucket.blob(file_path)

    filename = os.path.basename(file_path)
    local_path = os.path.join(temp_dir, filename)

    blob.download_to_filename(local_path)
    print(f"✅ Downloaded {filename} to {local_path}")

    return local_path, filename


def prepare_gemini_prompt() -> str:
    """
    Prepare the prompt for Gemini document analysis

    Returns:
        Formatted prompt string
    """

    prompt = """You are a specialized RFP Analyst. Your task is to extract structured metadata, pursuit details, production details, submission requirements, and a detailed compliance matrix from the provided RFP documents to initialize a project workspace.

INSTRUCTIONS:

1. **Project Metadata**: Carefully identify:
   - Project Name: Look for titles like "Request for Proposal for...", "RFP Title:", "Project:", or document headers
   - Issuer Name: Organization, agency, or company issuing the RFP (look for headers, letterheads, contact information)
   - Due Date: Submission deadline, proposal due date, closing date (convert to ISO 8601 format YYYY-MM-DDTHH:MM:SSZ). Look for phrases like "Due Date:", "Deadline:", "Closing Date:", "Submit by:"

2. **Pursuit Details**: Extract customer/issuer contact and approval chain information:
   - Customer Address: Physical address of the issuing organization (street, city, state, zip, country)
   - Contact Info: Primary contact person for RFP questions (name, email, phone, title/role)
   - Final Approver: Person with final approval authority (name, title, email if available)
   - Signer: Authorized contract signer (name, title, email if available)
   - Look for sections like "Point of Contact", "Issuing Officer", "Contracting Officer", "Contact Information", "Questions should be directed to"

3. **Production Details**: Extract submission format and delivery requirements:
   - Submission Format: Digital/electronic, physical/printed, or both
   - File Requirements: Allowed formats (PDF, Word, etc.), file size limits, naming conventions
   - Print Requirements: Number of copies, binding requirements, paper size
   - Delivery Method: Email, portal upload, mail, hand delivery
   - Mailing/Delivery Address: Where physical submissions should be sent (if different from customer address)
   - Special Instructions: Any specific formatting, packaging, or labeling requirements
   - Look for sections like "Submission Instructions", "Proposal Format", "Delivery Requirements", "How to Submit"

4. **Submission Requirements**: Extract ALL items that proposers must submit. Look for sections like:
   - "Submission Requirements", "Required Documents", "Proposal Components", "Deliverables", "What to Submit"
   - Forms (e.g., "Conflict of Interest Form", "Proposal Cover Sheet", "Budget Template", "Bid Form")
   - Documents (e.g., "Technical Proposal", "Financial Proposal", "Company Profile", "Executive Summary")
   - Certifications (e.g., "Insurance Certificate", "Business License", "Tax Clearance")
   - Attachments (e.g., "Work Samples", "References", "Resumes", "Past Performance")
   - Information to provide (e.g., "Project Timeline", "Pricing Structure", "Methodology")

5. **Compliance Matrix Extraction**: For EVERY specific mandate or requirement statement, populate the compliance_matrix array:
   - Look for "shall", "must", "will", "required to", "contractor will", "vendor shall" statements
   - Extract the EXACT text of each requirement
   - Identify the source section reference (e.g., "L.4.2.1", "Section 3.2", "Paragraph 5.1.3")
   - Note the page number where found
   - Map each requirement to ONE of these categories:
     * CERTIFICATION - Requires certifications, licenses, or accreditations
     * EXPERIENCE - Past performance, years of experience, similar projects
     * PERSONNEL - Staffing requirements, key personnel, qualifications
     * FORMAT - Document format, page limits, font requirements
     * SUBMISSION - Submission process, deadlines, delivery methods
     * FINANCIAL - Pricing, budget, financial capacity requirements
     * LEGAL - Legal compliance, insurance, bonding, contractual terms
     * TECHNICAL - Technical specifications, system requirements, capabilities
     * OTHER - Anything that doesn't fit above categories
   - Generate sequential IDs: CM-001, CM-002, CM-003, etc.

6. **De-duplication**: If a requirement appears in multiple files/sections, create ONE entry with multiple mentions in the mentions array

7. **Naming**: Use clear, task-friendly names:
   - Good: "Submit Technical Proposal", "Provide Insurance Certificate", "Complete Budget Form"
   - Bad: "Technical Proposal Document Submission Requirements Section 4.1"

8. **Required vs Optional**:
   - Set is_required=true if document explicitly states "required", "mandatory", "must", "shall"
   - Set is_required=false if document states "optional", "if applicable", "may"

9. **Source Location**: Be specific about where you found the requirement:
   - Good: "Section 4.1 - Submission Requirements, Page 12"
   - Bad: "In the document"

10. **Source Text**: For each requirement/item, include the EXACT verbatim text from the document:
    - Copy the relevant sentence or paragraph word-for-word
    - Keep it concise (max 500 characters)
    - This helps users verify the extraction against the original document

11. **Confidence Score**:
    - "high": Clearly stated requirement with explicit details
    - "medium": Implied requirement or unclear details
    - null: If very uncertain

12. **Handle Multiple Documents**: If multiple documents are provided, analyze ALL of them and aggregate findings

EXAMPLE OUTPUT (this is what your response should look like):
{
  "project_metadata": {
    "project_name": "Cloud Infrastructure Services RFP",
    "issuer_name": "Department of Technology",
    "due_date": "2024-03-15T17:00:00Z"
  },
  "pursuit_details": {
    "customer_address": {
      "street": "123 Enterprise Blvd",
      "city": "San Francisco",
      "state": "CA",
      "zip": "94105",
      "country": "USA"
    },
    "contact_info": {
      "name": "David Thompson",
      "title": "Procurement Officer",
      "email": "david.t@techcorp.com",
      "phone": "(555) 123-4567"
    },
    "final_approver": {
      "name": "Jennifer Walsh",
      "title": "CTO",
      "email": "j.walsh@techcorp.com"
    },
    "signer": {
      "name": "Robert Kim",
      "title": "Legal Director",
      "email": "r.kim@techcorp.com"
    },
    "source": {
      "source_file": "RFP_Main_Document.pdf",
      "source_location": "Cover Page and Section 1.3 - Contact Information",
      "confidence_score": "high"
    }
  },
  "production_details": {
    "submission_format": "Both digital and physical submissions required",
    "file_requirements": {
      "formats": ["PDF"],
      "max_file_size": "50MB",
      "naming_convention": "CompanyName_TechnicalProposal_YYYYMMDD.pdf"
    },
    "print_requirements": {
      "copies": 3,
      "binding": "3-ring binder",
      "paper_size": "Letter (8.5 x 11)",
      "additional_notes": "Printed copies for review committee"
    },
    "delivery_method": {
      "electronic": {
        "method": "Email",
        "destination": "proposals@techcorp.com",
        "portal_url": null
      },
      "physical": {
        "method": "Mail or Hand Delivery",
        "address": {
          "attention": "Procurement Office - RFP #2024-001",
          "street": "123 Enterprise Blvd, Suite 400",
          "city": "San Francisco",
          "state": "CA",
          "zip": "94105"
        }
      }
    },
    "special_instructions": "Mark all packages 'RFP Response - Do Not Open'. Electronic submissions must be received by 5:00 PM PST.",
    "source": {
      "source_file": "RFP_Main_Document.pdf",
      "source_location": "Section 3 - Submission Instructions, Pages 8-10",
      "confidence_score": "high"
    }
  },
  "submission_requirements": [
    {
      "response_item_name": "Submit Technical Proposal",
      "description": "Detailed technical approach including methodology, timeline, and deliverables. Must not exceed 20 pages and include system architecture diagrams.",
      "is_required": true,
      "mentions": [
        {
          "source_file": "RFP_Main_Document.pdf",
          "source_location": "Section 4.1 - Submission Requirements, Page 12",
          "source_text": "The Offeror shall submit a Technical Proposal not exceeding 20 pages that includes a detailed technical approach, methodology, timeline, deliverables, and system architecture diagrams.",
          "confidence_score": "high"
        }
      ]
    },
    {
      "response_item_name": "Provide Company Financial Statements",
      "description": "Audited financial statements for the past 3 years demonstrating financial stability and capacity to execute the project",
      "is_required": true,
      "mentions": [
        {
          "source_file": "RFP_Main_Document.pdf",
          "source_location": "Section 4.2 - Qualification Requirements, Page 15",
          "source_text": "Offerors must provide audited financial statements for the past three (3) fiscal years to demonstrate financial stability and capacity to execute the project.",
          "confidence_score": "high"
        },
        {
          "source_file": "Appendix_B.pdf",
          "source_location": "Page 2 - Required Documents Checklist, Item 7",
          "source_text": "Item 7: Three years of audited financial statements (required)",
          "confidence_score": "high"
        }
      ]
    },
    {
      "response_item_name": "Complete Conflict of Interest Form",
      "description": "Mandatory disclosure form regarding potential conflicts of interest. Form template provided in Appendix C.",
      "is_required": true,
      "mentions": [
        {
          "source_file": "RFP_Main_Document.pdf",
          "source_location": "Section 5.3 - Required Forms, Page 22",
          "source_text": "All offerors must complete and submit the Conflict of Interest Disclosure Form (Appendix C) with their proposal.",
          "confidence_score": "high"
        }
      ]
    }
  ],
  "compliance_matrix": [
    {
      "id": "CM-001",
      "requirement_text": "The Contractor shall provide 24/7 technical support with a maximum response time of 4 hours for critical issues.",
      "source_section": "L.4.2.1",
      "source_page": 15,
      "category": "TECHNICAL"
    },
    {
      "id": "CM-002",
      "requirement_text": "The Contractor must maintain ISO 27001 certification throughout the contract period.",
      "source_section": "L.5.1.3",
      "source_page": 18,
      "category": "CERTIFICATION"
    },
    {
      "id": "CM-003",
      "requirement_text": "The Contractor shall assign a dedicated Project Manager with at least 5 years of relevant experience.",
      "source_section": "L.6.2",
      "source_page": 21,
      "category": "PERSONNEL"
    },
    {
      "id": "CM-004",
      "requirement_text": "All proposals must be submitted in PDF format with a maximum file size of 25MB.",
      "source_section": "M.2.1",
      "source_page": 8,
      "category": "FORMAT"
    },
    {
      "id": "CM-005",
      "requirement_text": "The Contractor shall maintain professional liability insurance of at least $2 million.",
      "source_section": "L.8.4",
      "source_page": 25,
      "category": "LEGAL"
    }
  ]
}

OUTPUT FORMAT: Return ONLY the JSON object. No markdown code blocks, no conversational text, just the raw JSON.

Now analyze the provided document(s) and extract the information according to these instructions."""

    return prompt


def repair_truncated_json(json_text: str) -> Dict[str, Any]:
    """
    Attempt to repair truncated JSON by closing unclosed brackets and arrays.

    Args:
        json_text: Potentially truncated JSON string

    Returns:
        Parsed JSON dictionary with available data
    """
    import re

    # Track open brackets and braces
    open_braces = 0
    open_brackets = 0
    in_string = False
    escape_next = False

    for char in json_text:
        if escape_next:
            escape_next = False
            continue
        if char == '\\':
            escape_next = True
            continue
        if char == '"' and not escape_next:
            in_string = not in_string
            continue
        if in_string:
            continue
        if char == '{':
            open_braces += 1
        elif char == '}':
            open_braces -= 1
        elif char == '[':
            open_brackets += 1
        elif char == ']':
            open_brackets -= 1

    # Find a safe truncation point (end of last complete object/array)
    repaired = json_text.rstrip()

    # Remove trailing incomplete content
    # Look for common truncation patterns and clean them
    patterns_to_remove = [
        r',\s*$',  # Trailing comma
        r',\s*"[^"]*$',  # Incomplete key
        r':\s*"[^"]*$',  # Incomplete string value
        r':\s*$',  # Incomplete value
        r'"[^"]*$',  # Unterminated string
    ]

    for pattern in patterns_to_remove:
        repaired = re.sub(pattern, '', repaired)

    # Close any open structures
    repaired += ']' * open_brackets
    repaired += '}' * open_braces

    try:
        result = json.loads(repaired)
        print(f"✅ Successfully repaired truncated JSON")
        return result
    except json.JSONDecodeError:
        # If still failing, try a more aggressive approach - extract what we can
        print(f"⚠️ Could not fully repair JSON, extracting partial data...")
        return extract_partial_json(json_text)


def extract_partial_json(json_text: str) -> Dict[str, Any]:
    """
    Extract whatever valid JSON sections we can find from truncated text.

    Args:
        json_text: Truncated JSON string

    Returns:
        Dictionary with extracted sections
    """
    import re

    result = {
        'project_metadata': {'project_name': None, 'issuer_name': None, 'due_date': None},
        'pursuit_details': None,
        'production_details': None,
        'submission_requirements': [],
        'compliance_matrix': []
    }

    # Try to extract project_metadata
    metadata_match = re.search(r'"project_metadata"\s*:\s*(\{[^}]+\})', json_text)
    if metadata_match:
        try:
            result['project_metadata'] = json.loads(metadata_match.group(1))
        except:
            pass

    # Try to extract compliance_matrix items individually
    cm_items = re.findall(r'\{\s*"id"\s*:\s*"CM-\d+[^}]+\}', json_text)
    for item_str in cm_items:
        try:
            # Clean up the item string
            item_str = item_str.rstrip(',')
            if not item_str.endswith('}'):
                item_str += '}'
            item = json.loads(item_str)
            if item.get('id') and item.get('requirement_text'):
                result['compliance_matrix'].append(item)
        except:
            pass

    # Try to extract submission_requirements items
    sr_items = re.findall(r'\{\s*"response_item_name"\s*:[^}]+\}', json_text)
    for item_str in sr_items:
        try:
            item_str = item_str.rstrip(',')
            if not item_str.endswith('}'):
                item_str += '}'
            item = json.loads(item_str)
            if item.get('response_item_name'):
                result['submission_requirements'].append(item)
        except:
            pass

    print(f"📋 Extracted partial data: {len(result['compliance_matrix'])} compliance items, "
          f"{len(result['submission_requirements'])} submission requirements")

    return result


def call_gemini_for_shredding(files_data: List[tuple[bytes, str]]) -> Dict[str, Any]:
    """
    Call Gemini to extract metadata and submission requirements using multimodal capabilities

    Args:
        files_data: List of tuples (file_bytes, filename)

    Returns:
        Parsed JSON response from Gemini
    """
    try:
        # Initialize Gemini model
        # Using gemini-2.5-flash - higher output token limits
        model = GenerativeModel("gemini-2.5-flash")

        # Prepare parts for multimodal input
        parts = []

        # Add the prompt as the first part
        prompt = prepare_gemini_prompt()
        parts.append(Part.from_text(prompt))

        # Define MIME type mapping for Gemini-supported formats
        # Note: Gemini does NOT support .doc/.docx directly - they must be converted to text
        mime_mapping = {
            "pdf": "application/pdf",
            "txt": "text/plain",
            "csv": "text/csv",
            "json": "application/json",
            "png": "image/png",
            "jpg": "image/jpeg",
            "jpeg": "image/jpeg",
            "gif": "image/gif",
            "webp": "image/webp"
        }

        # File extensions that need text extraction (not supported by Gemini multimodal)
        text_extraction_extensions = {"doc", "docx"}

        # Add each file as a multimodal part
        for file_bytes, filename in files_data:
            file_ext = os.path.splitext(filename)[1].lower().lstrip('.')

            # Check if this file type needs text extraction
            if file_ext in text_extraction_extensions:
                # Convert .docx to text
                if file_ext == "docx":
                    try:
                        extracted_text = extract_text_from_docx(file_bytes)
                        # Add as text part with filename header
                        parts.append(Part.from_text(f"\n\n--- Document: {filename} ---\n\n{extracted_text}"))
                        print(f"✅ Converted {filename} to text ({len(extracted_text)} chars)")
                        continue
                    except Exception as e:
                        print(f"⚠️ Could not extract text from {filename}: {e}")
                        # Fall through to try binary upload (will likely fail)

                elif file_ext == "doc":
                    print(f"⚠️ .doc format not supported. Please convert {filename} to .docx or .pdf")
                    # Add a placeholder message
                    parts.append(Part.from_text(f"\n\n--- Document: {filename} (Unable to process .doc format - please convert to .docx or .pdf) ---\n\n"))
                    continue

            # For supported formats, use multimodal upload
            mime_type = mime_mapping.get(file_ext, "application/octet-stream")

            # Skip unsupported MIME types
            if mime_type == "application/octet-stream":
                print(f"⚠️ Unsupported file type: {filename} ({file_ext}). Skipping.")
                continue

            # Add file part
            file_part = Part.from_data(data=file_bytes, mime_type=mime_type)
            parts.append(file_part)

            # Add a text separator with filename
            parts.append(Part.from_text(f"\n\n--- End of document: {filename} ---\n\n"))

            print(f"✅ Added {filename} to multimodal request ({mime_type}, {len(file_bytes)} bytes)")

        # Safety settings to prevent blocking
        safety_settings = [
            SafetySetting(
                category="HARM_CATEGORY_HATE_SPEECH",
                threshold="BLOCK_NONE"
            ),
            SafetySetting(
                category="HARM_CATEGORY_DANGEROUS_CONTENT",
                threshold="BLOCK_NONE"
            ),
            SafetySetting(
                category="HARM_CATEGORY_SEXUALLY_EXPLICIT",
                threshold="BLOCK_NONE"
            ),
            SafetySetting(
                category="HARM_CATEGORY_HARASSMENT",
                threshold="BLOCK_NONE"
            ),
        ]

        # Generation config for JSON output
        # gemini-2.5-flash supports higher output token limits
        generation_config = {
            "temperature": 0.2,
            "max_output_tokens": 65536,  # Higher limit for gemini-2.5-flash
            "top_p": 0.8,
            "response_mime_type": "application/json",  # Ensures valid JSON output
        }

        print(f"🤖 Sending request to Gemini for document shredding ({len(files_data)} files)...")

        # Call Gemini with multimodal parts
        response = model.generate_content(
            parts,
            generation_config=generation_config,
            safety_settings=safety_settings,
            stream=False
        )

        print(f"✅ Received response from Gemini")

        # Parse JSON response
        response_text = response.text.strip()
        print(f"📝 Raw response length: {len(response_text)} characters")

        # Clean up response if needed (remove markdown code blocks)
        if response_text.startswith('```json'):
            response_text = response_text[7:]
        if response_text.startswith('```'):
            response_text = response_text[3:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
        response_text = response_text.strip()

        # Try to parse JSON, with recovery for truncated responses
        try:
            result = json.loads(response_text)
        except json.JSONDecodeError as json_err:
            print(f"⚠️ JSON parse error: {json_err}. Attempting to repair truncated JSON...")
            result = repair_truncated_json(response_text)

        print(f"✅ Successfully parsed JSON response")

        return result

    except Exception as e:
        print(f"❌ Error calling Gemini: {e}")
        import traceback
        traceback.print_exc()
        raise


def shred_documents(files: List[Dict[str, str]], org_id: str) -> Dict[str, Any]:
    """
    Main function to shred RFP documents and extract metadata

    Args:
        files: List of file dictionaries with 'file_id', 'filename', 'gcs_url'
        org_id: Organization ID

    Returns:
        Dictionary with project_metadata and submission_requirements
    """

    print(f"📄 Starting document shredding for {len(files)} files...")

    # Create temporary directory for downloads
    temp_dir = tempfile.mkdtemp(prefix="rfp_shredding_")

    try:
        # Step 1: Download all files from GCS and read as bytes
        files_data = []

        for file_info in files:
            filename = file_info['filename']
            gcs_url = file_info['gcs_url']

            print(f"📥 Downloading {filename} from GCS...")

            try:
                local_path, _ = download_file_from_gcs(gcs_url, temp_dir)

                # Read file as bytes
                with open(local_path, 'rb') as f:
                    file_bytes = f.read()

                files_data.append((file_bytes, filename))
                print(f"✅ Prepared {filename} for multimodal upload ({len(file_bytes)} bytes)")

            except Exception as e:
                print(f"❌ Error processing {filename}: {e}")
                # Continue with other files
                continue

        if not files_data:
            raise Exception("No files could be processed")

        # Step 2: Call Gemini with all files using multimodal input
        result = call_gemini_for_shredding(files_data)

        # Step 3: Validate and return result
        if not result.get('project_metadata'):
            result['project_metadata'] = {
                'project_name': None,
                'issuer_name': None,
                'due_date': None
            }

        # Validate pursuit_details structure
        if not result.get('pursuit_details'):
            result['pursuit_details'] = {
                'customer_address': None,
                'contact_info': None,
                'final_approver': None,
                'signer': None,
                'source': None
            }
        else:
            # Ensure all sub-fields exist
            pursuit = result['pursuit_details']
            if not pursuit.get('customer_address'):
                pursuit['customer_address'] = None
            if not pursuit.get('contact_info'):
                pursuit['contact_info'] = None
            if not pursuit.get('final_approver'):
                pursuit['final_approver'] = None
            if not pursuit.get('signer'):
                pursuit['signer'] = None
            if not pursuit.get('source'):
                pursuit['source'] = None

        # Validate production_details structure
        if not result.get('production_details'):
            result['production_details'] = {
                'submission_format': None,
                'file_requirements': None,
                'print_requirements': None,
                'delivery_method': None,
                'special_instructions': None,
                'source': None
            }
        else:
            # Ensure all sub-fields exist
            production = result['production_details']
            if not production.get('submission_format'):
                production['submission_format'] = None
            if not production.get('file_requirements'):
                production['file_requirements'] = None
            if not production.get('print_requirements'):
                production['print_requirements'] = None
            if not production.get('delivery_method'):
                production['delivery_method'] = None
            if not production.get('special_instructions'):
                production['special_instructions'] = None
            if not production.get('source'):
                production['source'] = None

        if not result.get('submission_requirements'):
            result['submission_requirements'] = []

        # Validate compliance_matrix structure
        if not result.get('compliance_matrix'):
            result['compliance_matrix'] = []
        else:
            # Ensure each item has required fields
            valid_categories = {'CERTIFICATION', 'EXPERIENCE', 'PERSONNEL', 'FORMAT',
                              'SUBMISSION', 'FINANCIAL', 'LEGAL', 'TECHNICAL', 'OTHER'}
            for idx, item in enumerate(result['compliance_matrix']):
                # Ensure ID exists
                if not item.get('id'):
                    item['id'] = f"CM-{str(idx + 1).zfill(3)}"
                # Ensure category is valid
                if item.get('category', '').upper() not in valid_categories:
                    item['category'] = 'OTHER'
                else:
                    item['category'] = item['category'].upper()
                # Ensure required fields exist
                if not item.get('requirement_text'):
                    item['requirement_text'] = ''
                if not item.get('source_section'):
                    item['source_section'] = None
                if not item.get('source_page'):
                    item['source_page'] = None

        # Log extraction summary
        print(f"✅ Document shredding complete!")
        print(f"   - Project Name: {result['project_metadata'].get('project_name')}")
        print(f"   - Issuer: {result['project_metadata'].get('issuer_name')}")
        print(f"   - Due Date: {result['project_metadata'].get('due_date')}")

        # Log pursuit details
        pursuit = result.get('pursuit_details', {})
        contact = pursuit.get('contact_info') if pursuit else None
        print(f"   - Contact: {contact.get('name') if contact else 'Not found'}")
        print(f"   - Final Approver: {pursuit.get('final_approver', {}).get('name') if pursuit.get('final_approver') else 'Not found'}")
        print(f"   - Signer: {pursuit.get('signer', {}).get('name') if pursuit.get('signer') else 'Not found'}")

        # Log production details
        production = result.get('production_details', {})
        print(f"   - Submission Format: {production.get('submission_format') if production else 'Not found'}")

        print(f"   - Submission Requirements Found: {len(result['submission_requirements'])}")
        print(f"   - Compliance Matrix Items Found: {len(result['compliance_matrix'])}")

        return result

    finally:
        # Clean up temporary directory
        import shutil
        try:
            shutil.rmtree(temp_dir)
            print(f"🧹 Cleaned up temporary directory: {temp_dir}")
        except Exception as e:
            print(f"⚠️ Could not clean up temp directory: {e}")


def shred_documents_endpoint_handler(request_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Flask endpoint handler for document shredding

    Args:
        request_data: Request JSON data with 'files' and 'org_id'

    Returns:
        Response dictionary
    """
    try:
        # Validate request
        files = request_data.get('files', [])
        org_id = request_data.get('org_id')

        if not files:
            return {
                'success': False,
                'error': 'No files provided'
            }, 400

        if not org_id:
            return {
                'success': False,
                'error': 'Organization ID is required'
            }, 400

        # Validate each file has required fields
        for file_info in files:
            if not all(key in file_info for key in ['file_id', 'filename', 'gcs_url']):
                return {
                    'success': False,
                    'error': 'Each file must have file_id, filename, and gcs_url'
                }, 400

        # Perform document shredding
        result = shred_documents(files, org_id)

        return {
            'success': True,
            'project_metadata': result['project_metadata'],
            'pursuit_details': result.get('pursuit_details'),
            'production_details': result.get('production_details'),
            'submission_requirements': result['submission_requirements'],
            'compliance_matrix': result.get('compliance_matrix', [])
        }, 200

    except Exception as e:
        print(f"❌ Error in document shredding endpoint: {e}")
        import traceback
        traceback.print_exc()

        return {
            'success': False,
            'error': 'Internal server error',
            'details': str(e)
        }, 500
