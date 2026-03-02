"""
Base Agent Class for RFP Extraction Agents
All agents inherit from this and implement extract() method
Uses AWS Bedrock (Claude) for AI processing
"""

import os
import json
import time
import tempfile
from abc import ABC, abstractmethod
from typing import Dict, Any, List, Tuple, Optional
from datetime import datetime

from google.cloud import storage
from google.oauth2 import service_account

# Import Bedrock client
from bedrock_client import BedrockClaude, BEDROCK_AVAILABLE

# For docx conversion
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class BaseExtractionAgent(ABC):
    """Base class for all extraction agents - uses AWS Bedrock (Claude)"""

    # Agent configuration
    AGENT_TYPE: str = "BASE"
    MAX_ITEMS: Optional[int] = None
    MAX_RETRIES: int = 3
    RETRY_DELAY: int = 5  # seconds

    def __init__(self):
        # Initialize Bedrock Claude client
        self.claude = BedrockClaude()
        self.max_tokens = 32768  # Higher limit for complex RFPs
        self.temperature = 0.2

    @abstractmethod
    def get_prompt(self) -> str:
        """Return the extraction prompt for this agent"""
        pass

    @abstractmethod
    def validate_result(self, result: Dict) -> Dict:
        """Validate and clean the extraction result"""
        pass

    def extract(self, files: List[Dict[str, str]], org_id: str) -> Dict[str, Any]:
        """
        Run extraction on provided files

        Args:
            files: List of file dictionaries with 'file_id', 'filename', 'gcs_url'
            org_id: Organization ID

        Returns:
            Dictionary with extraction results
        """
        print(f"🤖 [{self.AGENT_TYPE}] Starting extraction for {len(files)} files...")

        temp_dir = tempfile.mkdtemp(prefix=f"rfp_{self.AGENT_TYPE.lower()}_")

        try:
            # Download files and prepare for Gemini
            files_data = self._download_files(files, temp_dir)

            if not files_data:
                return {
                    "success": False,
                    "error": "No files could be processed",
                    "agent_type": self.AGENT_TYPE,
                }

            # Call Claude with the agent's prompt
            result = self._call_claude(files_data)

            # Validate and clean the result
            validated_result = self.validate_result(result)

            return {
                "success": True,
                "agent_type": self.AGENT_TYPE,
                "extracted_at": datetime.utcnow().isoformat(),
                **validated_result
            }

        except Exception as e:
            print(f"❌ [{self.AGENT_TYPE}] Error during extraction: {e}")
            import traceback
            traceback.print_exc()
            return {
                "success": False,
                "error": str(e),
                "agent_type": self.AGENT_TYPE,
            }
        finally:
            # Cleanup temp directory
            import shutil
            try:
                shutil.rmtree(temp_dir)
            except:
                pass

    def _download_files(self, files: List[Dict[str, str]], temp_dir: str) -> List[Tuple[bytes, str]]:
        """Download files from GCS and return as bytes"""
        files_data = []

        for file_info in files:
            filename = file_info['filename']
            gcs_url = file_info['gcs_url']

            try:
                local_path = self._download_file_from_gcs(gcs_url, temp_dir)

                with open(local_path, 'rb') as f:
                    file_bytes = f.read()

                files_data.append((file_bytes, filename))
                print(f"✅ [{self.AGENT_TYPE}] Prepared {filename} ({len(file_bytes)} bytes)")

            except Exception as e:
                print(f"❌ [{self.AGENT_TYPE}] Error processing {filename}: {e}")
                continue

        return files_data

    def _download_file_from_gcs(self, gcs_url: str, temp_dir: str) -> str:
        """Download a file from GCS to temp directory"""
        if not gcs_url.startswith('gs://'):
            raise ValueError(f"Invalid GCS URL: {gcs_url}")

        parts = gcs_url.replace('gs://', '').split('/', 1)
        bucket_name = parts[0]
        file_path = parts[1] if len(parts) > 1 else ''

        gcs_client = self._get_gcs_client()
        bucket = gcs_client.bucket(bucket_name)
        blob = bucket.blob(file_path)

        filename = os.path.basename(file_path)
        local_path = os.path.join(temp_dir, filename)

        blob.download_to_filename(local_path)
        return local_path

    def _get_gcs_client(self):
        """Get GCS client with proper credential handling"""
        cred_path = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS", "fire.json")

        if cred_path.startswith("{"):
            cred_dict = json.loads(cred_path)
            credentials = service_account.Credentials.from_service_account_info(cred_dict)
            return storage.Client(credentials=credentials, project=os.environ.get("GOOGLE_CLOUD_PROJECT"))
        else:
            if os.path.exists(cred_path):
                credentials = service_account.Credentials.from_service_account_file(cred_path)
                return storage.Client(credentials=credentials)
            return storage.Client()

    def _call_claude(self, files_data: List[Tuple[bytes, str]]) -> Dict[str, Any]:
        """Call Claude API via AWS Bedrock with the agent's prompt and files"""
        prompt = self.get_prompt()

        print(f"🤖 [{self.AGENT_TYPE}] Sending request to Claude via Bedrock ({len(files_data)} files)...")

        try:
            # Use the Bedrock client's document handling
            result = self.claude.call_claude_with_documents(
                prompt=prompt,
                documents=files_data,
                max_tokens=self.max_tokens,
                temperature=self.temperature,
                response_format="json"
            )

            print(f"✅ [{self.AGENT_TYPE}] Successfully parsed response from Claude")
            return result

        except Exception as e:
            print(f"❌ [{self.AGENT_TYPE}] Claude API error: {e}")
            raise

    def _sanitize_json_string(self, json_text: str) -> str:
        """
        Sanitize control characters in JSON strings that cause parse errors.
        Handles unescaped newlines, tabs, and other control chars within string values.
        """
        import re

        result = []
        in_string = False
        escape_next = False
        i = 0

        while i < len(json_text):
            char = json_text[i]

            if escape_next:
                result.append(char)
                escape_next = False
                i += 1
                continue

            if char == '\\':
                result.append(char)
                escape_next = True
                i += 1
                continue

            if char == '"':
                in_string = not in_string
                result.append(char)
                i += 1
                continue

            # If inside a string, escape control characters
            if in_string:
                if char == '\n':
                    result.append('\\n')
                elif char == '\r':
                    result.append('\\r')
                elif char == '\t':
                    result.append('\\t')
                elif ord(char) < 32:  # Other control characters
                    result.append(f'\\u{ord(char):04x}')
                else:
                    result.append(char)
            else:
                result.append(char)

            i += 1

        return ''.join(result)

    def _extract_text_from_docx(self, file_bytes: bytes) -> str:
        """Extract text from DOCX file"""
        import io
        doc = DocxDocument(io.BytesIO(file_bytes))

        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n\n".join(text_parts)

    def _repair_truncated_json(self, json_text: str) -> Dict[str, Any]:
        """Attempt to repair truncated JSON with multiple strategies"""
        import re

        # Strategy 1: Find and close unclosed structures
        def try_close_structures(text: str) -> Optional[Dict]:
            # Track open brackets
            open_braces = 0
            open_brackets = 0
            in_string = False
            escape_next = False
            last_valid_pos = 0

            for i, char in enumerate(text):
                if escape_next:
                    escape_next = False
                    continue
                if char == '\\':
                    escape_next = True
                    continue
                if char == '"' and not escape_next:
                    in_string = not in_string
                    if not in_string:
                        last_valid_pos = i
                    continue
                if in_string:
                    continue
                if char in '{}[]':
                    last_valid_pos = i
                if char == '{':
                    open_braces += 1
                elif char == '}':
                    open_braces -= 1
                elif char == '[':
                    open_brackets += 1
                elif char == ']':
                    open_brackets -= 1

            # Close unclosed structures
            repaired = text.rstrip()

            # If we're inside a string, try to close it
            if in_string:
                # Find the last quote and truncate there, then close the string
                repaired = repaired.rstrip()
                # Remove partial string content
                while repaired and repaired[-1] != '"':
                    repaired = repaired[:-1]
                if not repaired.endswith('"'):
                    repaired += '"'

            # Remove trailing incomplete content (commas, colons, partial keys)
            while repaired:
                stripped = repaired.rstrip()
                if stripped and stripped[-1] in '"}]':
                    break
                repaired = repaired[:-1]

            # Recount after cleanup
            open_braces = 0
            open_brackets = 0
            in_string = False
            escape_next = False
            for char in repaired:
                if escape_next:
                    escape_next = False
                    continue
                if char == '\\':
                    escape_next = True
                    continue
                if char == '"' and not escape_next:
                    in_string = not in_string
                    continue
                if in_string:
                    continue
                if char == '{':
                    open_braces += 1
                elif char == '}':
                    open_braces -= 1
                elif char == '[':
                    open_brackets += 1
                elif char == ']':
                    open_brackets -= 1

            # Add closing brackets
            repaired += ']' * max(0, open_brackets)
            repaired += '}' * max(0, open_braces)

            try:
                return json.loads(repaired)
            except json.JSONDecodeError:
                return None

        # Strategy 2: Extract partial data using regex
        def extract_partial_data(text: str) -> Dict[str, Any]:
            result = {}

            # Try to extract each major section
            sections = [
                ('volume_structure', r'"volume_structure"\s*:\s*\[(.*?)\]', list),
                ('required_attachments', r'"required_attachments"\s*:\s*\[(.*?)\]', list),
                ('eligibility_items', r'"eligibility_items"\s*:\s*\[(.*?)\]', list),
                ('risks', r'"risks"\s*:\s*\[(.*?)\]', list),
                ('competitive_insights', r'"competitive_insights"\s*:\s*\[(.*?)\]', list),
                ('format_requirements', r'"format_requirements"\s*:\s*(\{[^}]+\})', dict),
                ('key_dates', r'"key_dates"\s*:\s*(\{[^}]+\})', dict),
                ('pricing_intelligence', r'"pricing_intelligence"\s*:\s*(\{[^}]+\})', dict),
                ('go_no_go_recommendation', r'"go_no_go_recommendation"\s*:\s*(\{[^}]+\})', dict),
            ]

            for key, pattern, expected_type in sections:
                try:
                    match = re.search(pattern, text, re.DOTALL)
                    if match:
                        content = match.group(1) if expected_type == dict else f'[{match.group(1)}]'
                        parsed = json.loads(content)
                        result[key] = parsed
                except:
                    if expected_type == list:
                        result[key] = []
                    else:
                        result[key] = {}

            return result

        # Try Strategy 1 first
        result = try_close_structures(json_text)
        if result:
            print(f"✅ [{self.AGENT_TYPE}] Repaired JSON using structure closure")
            return result

        # Try Strategy 2 as fallback
        print(f"⚠️ [{self.AGENT_TYPE}] Structure closure failed, extracting partial data...")
        result = extract_partial_data(json_text)
        if any(result.values()):
            print(f"✅ [{self.AGENT_TYPE}] Extracted partial data: {[k for k, v in result.items() if v]}")
            return result

        print(f"❌ [{self.AGENT_TYPE}] JSON repair failed completely")
        return {}
